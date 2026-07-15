import asyncio
import json
import os
from typing import Any, Dict, List, Optional
from pathlib import Path
from datetime import datetime
from typing import Dict, Any, Optional

import litellm
from litellm import acompletion
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pdfplumber

import config
import utils

logger = utils.setup_logging("generator")

CV_BASE_FILE = config.DATA_DIR / "cv_base.md"
CV_BASE_PDF = config.DATA_DIR / "cv_base.pdf"
LM_BASE_FILE = config.DATA_DIR / "lm_base.md"
LM_BASE_PDF = config.DATA_DIR / "lm_base.pdf"

PROMPT_CV = """Tu es un expert en recrutement et en rédaction de CV pour des postes en cybersécurité et en intelligence artificielle. Tu produis un CV professionnel (CDI/CDD) parfaitement CIBLÉ pour l'offre ci-dessous, à partir du CV de base du candidat.

## Règles
- CIBLAGE : sélectionne et place en avant EN PRIORITÉ les expériences, compétences, projets et certifications du CV de base qui collent le mieux à CETTE offre. Résume ou écarte ce qui est hors-sujet.
- MOTS-CLÉS / ATS : réutilise le vocabulaire de l'annonce (normes, outils, technologies, méthodologies) quand il correspond à des éléments RÉELS du candidat, pour passer les filtres automatiques.
- ACCROCHE : rédige un « Résumé » de 2-3 phrases adapté spécifiquement à ce poste et cette entreprise.
- ORDRE : classe sections et puces par pertinence décroissante pour l'offre.
- NE JAMAIS INVENTER : n'ajoute aucune expérience, compétence, diplôme, certification ou chiffre absent du CV de base. Si un prérequis de l'offre manque au candidat, ne le fabrique pas.
- CONCIS & FACTUEL : idéalement 1 page. Verbes d'action, résultats concrets quand ils existent dans le CV de base.
- IDENTITÉ : conserve telles quelles les vraies coordonnées du candidat (nom, email, téléphones, LinkedIn/GitHub).

## Format de sortie — Markdown STRICT (converti en DOCX ; utilise # pour le nom, ## pour les sections, - pour les puces)
# Prénom NOM
Ville · email · téléphone · LinkedIn · GitHub

## Résumé
[2-3 phrases ciblées sur CE poste]

## Compétences
- [regroupées par thème, les plus pertinentes d'abord]

## Expériences professionnelles
[les plus pertinentes, avec des puces orientées résultats]

## Formation

## Certifications

## Projets
[uniquement ceux qui servent cette candidature]

---

## OFFRE VISÉE
- Intitulé : {titre}
- Entreprise : {entreprise}
- Description de l'offre :
{description}

## CV DE BASE DU CANDIDAT (source unique de vérité — ne rien ajouter au-delà)
{cv_base}

Réponds UNIQUEMENT avec le CV final en Markdown, sans aucun commentaire ni texte autour."""


PROMPT_LM = """Tu es un expert en lettres de motivation. Génère une lettre de motivation personnalisée pour cette offre.

Format de sortie (Markdown):
# [Nom Prénom]
[Adresse]
[Téléphone] | [Email]

[Date]

Madame, Monsieur,

[Intro percutante - pourquoi cette entreprise et ce poste]

[Paragraphe 1 - Mon parcours et compétences]
[Paragraphe 2 - Ma motivation pour cette entreprise]
[Paragraphe 3 - Mon apport potentiel]

Cordialement,
[Nom Prénom]

---

offre:
- Titre: {titre}
- Entreprise: {entreprise}
- Description: {description}

Mon CV de base:
{cv_base}

LM originale:
{lm_base}

Réponds uniquement avec la lettre en Markdown."""


async def read_base_files() -> Dict[str, str]:
    files = {}

    for name, path in [("cv", CV_BASE_FILE), ("lm", LM_BASE_FILE)]:
        if path.exists():
            with open(path, "r", encoding="utf-8") as f:
                files[name] = f.read()
        else:
            logger.warning(f"Base file not found: {path}")
            files[name] = ""

    return files


async def extract_text_from_pdf(pdf_path: Path) -> str:
    if not pdf_path.exists():
        return ""

    text_parts = []
    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                text_parts.append(text)

    return "\n\n".join(text_parts)


def _api_key_for(model: str) -> Optional[str]:
    """Résout la clé API à passer à litellm selon le préfixe du modèle."""
    prefixes = {
        "openrouter/": "OPENROUTER_API_KEY",
        "gemini/": "GEMINI_API_KEY_1",
        "groq/": "GROQ_API_KEY_1",
        "nvidia_nim/": "NVIDIA_NIM_API_KEY",
        "cerebras/": "CEREBRAS_API_KEY",
        "mistral/": "MISTRAL_API_KEY",
    }
    for pref, env in prefixes.items():
        if model.startswith(pref):
            return os.getenv(env)
    return None


async def _generate_markdown(system: str, prompt: str, max_tokens: int = 6000) -> Optional[str]:
    """Essaie le modèle de raisonnement primaire puis le fallback (config).
    Nettoie le raisonnement (<think>...</think>) et renvoie le Markdown final."""
    models = [config.GENERATOR_PRIMARY_MODEL, config.GENERATOR_FALLBACK_MODEL]
    for model in models:
        try:
            response = await acompletion(
                model=model,
                api_key=_api_key_for(model),
                messages=[
                    {"role": "system", "content": system},
                    {"role": "user", "content": prompt},
                ],
                max_tokens=max_tokens,
            )
            content = utils.strip_reasoning(response.choices[0].message.content)
            if content:
                return content
            logger.warning(f"Génération: réponse vide ({model}), essai suivant")
        except Exception as e:
            logger.warning(f"Génération échouée ({model}): {e}")
    logger.error("Génération: tous les modèles ont échoué")
    return None


async def generate_cv_content(
    job: Dict[str, Any],
    cv_base: str,
    lm_base: str
) -> Optional[str]:
    prompt = PROMPT_CV.format(
        titre=job.get("titre", ""),
        entreprise=job.get("entreprise", ""),
        description=(job.get("description", "") or "Non fournie")[:4000],
        cv_base=cv_base,  # CV complet (plus de troncature)
    )
    return await _generate_markdown(
        "Tu es un expert en rédaction de CV pour la cybersécurité et l'IA. "
        "Réponds uniquement avec le CV en Markdown, sans texte autour.",
        prompt,
        max_tokens=8000,
    )


async def generate_lm_content(
    job: Dict[str, Any],
    cv_base: str,
    lm_base: str
) -> Optional[str]:
    prompt = PROMPT_LM.format(
        titre=job.get("titre", ""),
        entreprise=job.get("entreprise", ""),
        description=(job.get("description", "") or "Non fournie")[:4000],
        cv_base=cv_base,       # CV complet
        lm_base=lm_base[:1500],
    )
    return await _generate_markdown(
        "Tu es un expert en lettres de motivation. Réponds uniquement en Markdown.",
        prompt,
        max_tokens=6000,
    )


def parse_markdown_cv(md_text: str) -> Dict[str, Any]:
    sections = {}

    current_section = None
    current_content = []

    for line in md_text.split("\n"):
        line = line.strip()

        if line.startswith("## "):
            if current_section:
                sections[current_section] = "\n".join(current_content).strip()
            current_section = line[3:].strip()
            current_content = []
        elif line.startswith("# ") and not line.startswith("##"):
            sections["header"] = line[2:].strip()
        elif current_section:
            current_content.append(line)

    if current_section:
        sections[current_section] = "\n".join(current_content).strip()

    return sections


def create_docx_from_markdown(md_text: str, output_path: Path) -> bool:
    doc = Document()

    for line in md_text.split("\n"):
        line = line.strip()

        if not line:
            continue

        if line.startswith("# "):
            heading = doc.add_heading(line[2:], level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)

        elif line.startswith("- "):
            doc.add_paragraph(line, style="List Bullet")

        elif line.startswith("**") and line.endswith("**"):
            p = doc.add_paragraph()
            p.add_run(line.replace("**", "")).bold = True

        else:
            p = doc.add_paragraph(line)

    doc.save(output_path)
    logger.info(f"DOCX saved: {output_path}")
    return True


async def generate_for_job(job: Dict[str, Any]) -> Optional[Dict[str, Path]]:
    """CV ciblé via le module cvgen (sélection déterministe, PDF 1 page).
    La lettre de motivation garde l'ancien circuit LLM->DOCX (non prioritaire) ;
    son échec ne bloque jamais le CV."""
    entreprise = utils.safe_filename(job.get("entreprise", "entreprise_unknown"))
    date_str = datetime.now().strftime("%Y%m%d")

    output_dir = config.GENERATED_DIR / f"{entreprise}_{date_str}"
    output_dir.mkdir(parents=True, exist_ok=True)

    logger.info(f"Generating CV/LM for {entreprise}...")

    from cvgen import generate_cv_for_job

    try:
        # Profil 100 % déterministe (variante taggée + disponibilité).
        # Pour réactiver l'accroche LLM : hook=cvgen.accroche.llm_accroche.
        html_path, pdf_path, country = generate_cv_for_job(job, output_dir)
        cv_path = pdf_path or html_path  # sans moteur PDF, le HTML reste livrable
        logger.info(f"CV ({country}): {cv_path}")
    except Exception as e:
        logger.error(f"Failed to generate CV for {entreprise}: {e}")
        return None

    # --- Note d'analyse d'écarts offre <-> CV (LLM, best-effort) ---
    try:
        from cvgen.gap_analysis import write_gap_notes
        notes = write_gap_notes(job, html_path, output_dir)
        if notes:
            logger.info(f"Notes d'analyse : {notes}")
    except Exception as e:
        logger.warning(f"Analyse d'écarts non générée pour {entreprise}: {e}")

    # --- Lettre de motivation : ancien circuit, best-effort ---
    lm_ok = False
    lm_path = output_dir / "lm_personnalisee.docx"
    try:
        base_files = await read_base_files()
        if base_files.get("cv"):
            lm_content = await generate_lm_content(
                job, base_files.get("cv", ""), base_files.get("lm", "")
            )
            if lm_content:
                create_docx_from_markdown(lm_content, lm_path)
                lm_ok = True
    except Exception as e:
        logger.warning(f"LM non générée pour {entreprise}: {e}")

    return {
        "cv": cv_path,
        "lm": lm_path if lm_ok else None,
        "dir": output_dir,
    }


async def generate_all(
    jobs: List[Dict[str, Any]]
) -> Dict[str, List[Dict[str, Any]]]:
    generated = []

    for job in jobs:
        logger.info(f"Generating for: {job.get('entreprise', '')}")

        result = await generate_for_job(job)
        if result:
            generated.append({
                **job,
                "files": result,
            })

        await asyncio.sleep(1)

    return generated